'''
Generate a report aggregating all the diagnostics figures together.
Uses python-docx to generate a docx file
'''
import pandas
import docx
from docx.shared import Inches

import longitudinal_statistics

def generate_report(top_phenotypes, phecode_info, OUTDIR):
    doc = docx.Document()
    doc.add_heading("Cox Proportional Hazards Diagnostics")
    doc.add_paragraph("""This document contains supplemental tests performed to asses the impact of various assumptions in the Cox proportional hazards models. These assumptions are the proportional hazards assumption (that hazard ratios are constant over time) and the linearity of the effect. We also assess a competing outcomes models, where death is considered as a competing outcome rather than as a censoring.""")
    doc.add_paragraph("This document is large and we recommend navigating by the headings (enable View -> Navigation Pane if using MS Word or similar in other document viewers).")

    # ZPH
    doc.add_heading("Proportional Hazards Assumption", level=2)
    doc.add_paragraph("We assess the proprtional hazards assumption through the use of the R package 'survival' and it's function called 'cox.zph'. This examines the Schoenfeld residuals. The Schoenfeld residuals 'can essentially be thought of as the observed minus the expected values of the covariates at each failure time' (Steffensmeier & Jones, 2004: p 121).")
    doc.add_paragraph("Below, we plot for a selection of our top phenotypes, the Schoenfeld residual for each particiapnt in the analysis in black, against the time (in years) since the actigraphy measurement was taken. In red, a fit curve for the residuals. If the proportional hazards model is violated, then this curve will be non-horizontal. This is done for each of the covariates in the model.")
    doc.add_paragraph("For each, we also give the p-value from cox.zph of a deviation from horizontal. Due to the high number of subjects, we consider a p < 0.01 threshold for significance and add a star (*) to those p-values.")

    zph_ps = pandas.read_csv(OUTDIR / "diagnosis_figs" / "zph.ps.txt", sep="\t",index_col=0).T
    for phecode in top_phenotypes:
        phenotype = phecode_info.phenotype[phecode]
        doc.add_heading(f"{phenotype}", level=3)
        for covariate in longitudinal_statistics.COVARIATES:
            diagnostic_plot = OUTDIR / "diagnosis_figs" / f"{phecode}.{covariate}.png"
            if diagnostic_plot.exists():
                p = zph_ps.loc[phecode, covariate]
                par = doc.add_paragraph(f"{phenotype} - {covariate} - ")
                par.add_run(f"p = {p:0.3f}{'*' if p < 0.01 else ''}").bold = (p < 0.01)
                doc.add_picture(str(diagnostic_plot), width=Inches(3))
            else:
                print(f"Could not find file {diagnostic_plot}")

    # Time-varying
    doc.add_page_break()
    doc.add_heading("Time-varying Effects", level=2)
    doc.add_paragraph("Based off the identified significant deviations from the Cox proportional hazards model, we next introduce models that include time-varying effects for all the identified covariates that violate the assumptions. This corrects for the deviations from the constant hazard ratios assumption.")
    doc.add_paragraph("The results of running each of these models is displayed below. The table displays the p-value and effect size of the temperature amplitude effect. These can be compared to the main results in the manuscript for inconsistent results.")
    time_varying_results = pandas.read_csv(OUTDIR / "diagnosis_figs" / "time_varying_amp_summary.txt", sep="\t", index_col=0).T
    table = doc.add_table(rows=1, cols=time_varying_results.shape[1]+1)
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(time_varying_results.columns):
        hdr_cells[i+1].text = c
    for phecode, res in time_varying_results.iterrows():
        row = table.add_row().cells
        row[0].text = phecode_info.phenotype[phecode]
        for i, c in enumerate(time_varying_results.columns):
            row[i+1].text = f"{res[c]:0.4e}"

    # Nonlinearity
    doc.add_page_break()
    doc.add_heading("Non-linear Effects", level=2)
    doc.add_paragraph("We next assess whether the temperature amplitude variable may have a non-linear relation with the hazard ratio for each of the phenotypes. To do this, we use a spline fit (specifically, a pspline with df=3).")
    doc.add_paragraph("Below, we plot each nonlinear model by the HR versus temperature amplitude along with the confidence interval. We note apparent non-linearities that occur outside of the typical range of temperature amplitude values (0.62-4.2 degrees C for the 2.5th-97.5th percentiles) are generally unimportant and nonsignificant.")

    for phecode in top_phenotypes:
        phenotype = phecode_info.phenotype[phecode]
        nonlinear_plot = OUTDIR / "diagnosis_figs" / "nonlinear"/ f"{phecode}.nonlinear.temp_amplitude.png"
        if nonlinear_plot.exists():
            doc.add_heading(f"{phenotype}", level=3)
            doc.add_picture(str(nonlinear_plot), width=Inches(3))
        else:
            print(f"Could not find file {nonlinear_plot}")

    # Competing Outcomes
    doc.add_page_break()
    doc.add_heading("Competing outcomes", level=2)
    doc.add_paragraph("Finally, we consider whether including death as a competing outcome influences the model. In our main manuscript, we consider death to be a censoring event, equivalent to the end of data collection for that individual. Here we instead simultaneously model both death and diagnosis of the phenotype as potential outcomes.")
    doc.add_paragraph("The results of running each of these models is displayed below. The table displays the p-value and effect size of the temperature amplitude effect. These can be compared to the main results in the manuscript for inconsistent results.")
    comp_outs = pandas.read_csv(OUTDIR / "diagnosis_figs" / "competing_outcomes_amp_summary.txt", sep="\t", index_col=0).T
    table = doc.add_table(rows=1, cols=comp_outs.shape[1]+1)
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(comp_outs.columns):
        hdr_cells[i+1].text = c
    for phecode, res in comp_outs.iterrows():
        row = table.add_row().cells
        row[0].text = phecode_info.phenotype[phecode]
        for i, c in enumerate(comp_outs.columns):
            row[i+1].text = f"{res[c]:0.4e}"

    # Output
    doc.save(str(OUTDIR / "diagnostics_report.docx"))